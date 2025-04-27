import html
import os
from lxml import html as xhtml
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from bs4 import BeautifulSoup
import pdfkit
import pdfplumber
from docx import Document
from docx.shared import RGBColor
from reportlab import rl_config
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle
from pdf2docx import Converter as Conv
from docx2pdf import convert
from project.CellS import CellS
from reportlab.pdfbase.pdfmetrics import stringWidth

class Convertor:
	@staticmethod
	def convert_to_pdf(file_path, mode, file_name="html_to_pdf.pdf"):
		if mode == 1:
			doc = Document(file_path)
			table = doc.tables[0]
			rows = len(table.rows)
			columns = len(table.columns)
			list_cells = []
			texts = []
			for row_idx, row in enumerate(table.rows):
				row_cells = []  # 存储当前行的所有单元格对象
				text = []  # 存储当前行的所有单元格文本
				for cell_idx, cell in enumerate(row.cells):
					cells = Convertor.check_cell_style(cell)
					cells.text = cell.text.strip()  # 设置单元格的文本内容（去除首尾空格）
					# 记录单元格的行列位置
					cells.row = row_idx + 1
					cells.col = cell_idx + 1
					# cells.colSpan = is_grid_span(cell)
					row_cells.append(cells)  # 所有行的单元格对象
					text.append(cell.text.strip())
				# 存储提取到的文本和该行单元格信息
				list_cells.append(row_cells)
				texts.append(text)
			# print(texts)
			Convertor.pdf_convertor(list_cells, texts, columns, rows)
			return True
		elif mode == 2:
			try:
				path_to_wkhtmltopdf = r".\wkhtmltox\bin\wkhtmltopdf.exe"
				config = pdfkit.configuration(wkhtmltopdf=path_to_wkhtmltopdf)
				folder_path = r"./output_pdf"
				os.makedirs(folder_path, exist_ok=True)
				full_path = os.path.join(folder_path, file_name)
				pdfkit.from_file(file_path, full_path, configuration=config, options={"--encoding": "utf8"})
				return True
			except Exception as e:
				return False

	@staticmethod
	def check_cell_style(cell):
		"""检测单元格"""
		# 获得该单元格的段落
		cells = CellS()
		for paragraph in cell.paragraphs:
			# 获得该段落中的文本块
			for run in paragraph.runs:
				if run.font.name is None:
					cells.name = 'Calibri'
				else:
					cells.name = run.font.name
					print(run.font.name)
				if run.font.size is None:
					cells.size = 11
				else:
					cells.size = run.font.size.pt
				# 检查粗体和斜体
				if run.font.bold:
					cells.bold = True
				if run.font.italic:
					cells.italic = True
				# 检查颜色（确保颜色已显式设置）
				if run.font.color and run.font.color.rgb:
					# font_colors.add(run.font.color.rgb)
					cells.color = Convertor.rgb_to_hex(run.font.color.rgb)
		return cells

	@staticmethod
	def hex_to_rgb(hex_color):
		# 去掉 # 前缀（如果存在）
		hex_color = hex_color.lstrip('#')

		# 提取红、绿、蓝部分
		red = hex_color[0:2]
		green = hex_color[2:4]
		blue = hex_color[4:6]

		# 转换为十进制
		r = float(int(red, 16) / 255.0)
		g = float(int(green, 16) / 255.0)
		b = float(int(blue, 16) / 255.0)
		return r, g, b

	@staticmethod
	def rgb_to_hex(rgb_color):
		"""将 RGBColor 转换为十六进制字符串"""
		if not isinstance(rgb_color, RGBColor):
			return None
		rgb_values = str(rgb_color)
		print(rgb_values)
		return str('#' + rgb_values)

	# 从 RGBColor 对象中提取 RGB 值

	@staticmethod
	def register_font():
		rl_config.autoGenerateTTFName = True  # 注意正确参数名称

		font_folder = r"..\word"  # 设置字体文件夹路径
		ttf_files = [f for f in os.listdir(font_folder) if f.endswith('.ttf')]

		for ttf_file in ttf_files:
			font_path = os.path.join(font_folder, ttf_file)
			try:
				# 直接使用文件名作为字体名称的备选方案
				safe_name = os.path.splitext(ttf_file)[0]
				# 清理非法字符（关键步骤！）
				safe_name = "".join(c for c in safe_name if c.isalnum() or c in ('-', '_'))
				safe_name = safe_name[:30]  # 限制长度
				# 尝试注册字体
				pdfmetrics.registerFont(TTFont(safe_name, font_path))
			# print(f"字体 '{safe_name}' 注册成功")
			except Exception as e:
				pass  # 截断错误信息避免过长

	@staticmethod
	def pdf_convertor(list_cells, texts, column, row, output_filename="word_to_pdf.pdf"):

		Convertor.register_font()
		# 获取所有已注册的字体名称
		# argW = []
		# argH = []
		# # 创建PDF文档的行高和行宽列表
		# for i in range(column):
		# 	argW.append(80)
		# for i in range(row):
		# 	argH.append(40)
		font_name = "simhei"
		font_size = 11
		padding = 4
		col_widths = []
		# 获取字符串宽度函数

		# 遍历每列
		for col_idx in range(len(texts[0])):
			max_width = 0
			# 遍历每行该列的内容
			for row in texts:
				cell_content = str(row[col_idx])
				# 计算文字宽度（考虑换行符）
				lines = cell_content.split('\n')
				for line in lines:
					width = stringWidth(line, font_name, font_size)
					if width > max_width:
						max_width = width
			# 列宽 = 最大文字宽度 + 左右边距
			col_widths.append(max_width + 2 * padding)

		# ------------------------- 计算行高 -------------------------
		row_heights = []
		line_height = font_size * 1.2  # 单行高度（含行距）
		for row in texts:
			max_lines = 1
			for cell in row:
				cell_content = str(cell)
				lines = cell_content.split('\n')
				# 根据列宽计算自动换行后的行数
				for col_idx, content in enumerate(cell):
					lines = []
					current_line = []
					current_width = 0
					for word in content.split():
						word_width = stringWidth(word, font_name, font_size)
						if current_width + word_width > col_widths[col_idx]:
							lines.append(' '.join(current_line))
							current_line = [word]
							current_width = word_width
						else:
							current_line.append(word)
							current_width += word_width + stringWidth(' ', font_name, font_size)
					lines.append(' '.join(current_line))
					line_count = len(lines)
					if line_count > max_lines:
						max_lines = line_count
			# 行高 = 行数 * 单行高度 + 上下边距
			row_heights.append(max_lines * line_height + 2 * padding)

		print("*******************")
		print(row_heights)
		print(col_widths)
		print("*******************")
		folder_path = r"./output_pdf"
		os.makedirs(folder_path, exist_ok=True)
		full_path = os.path.join(folder_path, output_filename)

		doc = SimpleDocTemplate(full_path, pagesize=A4)
		table = Table(texts, colWidths=col_widths, rowHeights=row_heights)
		style = TableStyle([
			# 表格边框
			('GRID', (0, 0), (-1, -1), 1, colors.black),  # 全表格边框
			# 对齐方式
			('ALIGN', (0, 0), (-1, -1), 'CENTER'),  # 居中对齐
			('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),  # 垂直居中
			('FONT', (0, 0), (-1, -1), "AlibabaPuHuiTi-3-45-Light"),
			('FONTSIZE', (0, 0), (- 1, - 1), 12)
		])

		for row_cells in list_cells:
			for cells in row_cells:
				Convertor.cellStyle(style, cells)
		table.setStyle(style)
		table.hAlign = 'CENTER'  # 设置表格在页面中水平居中
		# 构建PDF文档
		elements = [table]
		doc.build(elements)

	@staticmethod
	def cellStyle(style, cells):
		if cells.color:
			(r, g, b) = Convertor.hex_to_rgb(cells.color)
			style.add('TEXTCOLOR', (cells.row - 1, cells.col - 1), (cells.row - 1, cells.col - 1), (r, g, b))

	@staticmethod
	def convertToHtml(list_cells):
		html_lines = ['<!DOCTYPE html>']
		html_lines.append('<html lang="zh-CN">')
		html_lines.append('<head>')
		html_lines.append('<meta charset="UTF-8">')
		html_lines.append('<meta name="viewport" content="width=device-width, initial-scale=1.0">')
		html_lines.append('<title>优雅表格</title>')
		html_lines.append('<style>')
		# 全局样式
		html_lines.append('''
		       body {
		           font-family: 'Segoe UI', system-ui, sans-serif;
		           background: #f0f2f5;
		           margin: 0;
		           min-height: 100vh;
		           display: flex;
		           justify-content: center;
		           align-items: center;
		       }

		       .table-container {
		           background: white;
		           border-radius: 12px;
		           box-shadow: 0 8px 24px rgba(0,0,0,0.05);
		           padding: 24px;
		           overflow-x: auto;
		           max-width: 90vw;
		       }

		       .data-table {
		           border-collapse: collapse;
		           width: auto;
		           min-width: 600px;
		       }

		       .data-table td {
		           padding: 14px 20px;
		           border: 1px solid #e4e7ed;
		           transition: all 0.2s ease;
		           position: relative;
		           color: #606266;
		           font-size: 14px;
		           line-height: 1.5;
		       }

		       .data-table tr:nth-child(even) td {
		           background-color: #f8f9fa;
		       }

		       .data-table tr:hover td {
		           background-color: #f5f7fa;
		           box-shadow: 0 2px 8px rgba(0,0,0,0.05);
		       }

		       /* 动态样式 */
		       .text-bold { font-weight: 600; }
		       .text-italic { font-style: italic; }
		       .text-large { font-size: 16px; }
		       .text-xlarge { font-size: 18px; }
		   ''')
		html_lines.append('</style>')
		html_lines.append('</head>')
		html_lines.append('<body>')
		html_lines.append('<div class="table-container">')
		html_lines.append('<table class="data-table">')

		for row_cells in list_cells:
			html_lines.append('<tr>')
			for cell in row_cells:
				# 动态生成类名
				classes = []
				if cell.bold:
					classes.append('text-bold')
				if cell.italic:
					classes.append('text-italic')
				if cell.size:
					if cell.size > 14:
						classes.append('text-xlarge' if cell.size > 16 else 'text-large')

				# 构建单元格属性
				attrs = []
				if classes:
					attrs.append(f'class="{" ".join(classes)}"')
				if cell.color:
					attrs.append(f'style="color: {cell.color}"')

				# 转义文本内容
				text = html.escape(cell.text) if cell.text is not None else '&nbsp;'

				html_lines.append(f'<td {" ".join(attrs)}>{text}</td>')
			html_lines.append('</tr>')

		html_lines.append('</table>')
		html_lines.append('</div>')
		html_lines.append('</body>')
		html_lines.append('</html>')
		return '\n'.join(html_lines)

	@staticmethod
	def convert_to_word(file_path, mode, file_name="to_word.docx"):
		if mode == 1:
			cv = Conv(file_path)
			folder_path = r"./output_word"
			os.makedirs(folder_path, exist_ok=True)
			full_path = os.path.join(folder_path, file_name)
			cv.convert(full_path)
			cv.close()
			return True
		elif mode == 2:
			try:
				with open(file_path, "r", encoding="utf-8") as html_file:
					html_content = html_file.read()
				# 使用lxml解析HTML
				tree = xhtml.fromstring(html_content)
				# 提取所有表格
				tables = tree.xpath('//table')
				# 如果有多个表格，可以选择第一个或全部
				if tables:
					# 提取第一个表格
					table = tables[0]
					# 将表格内容转换为字符串
					table_html = xhtml.tostring(table, encoding='unicode')
					# 使用BeautifulSoup解析HTML表格
					soup = BeautifulSoup(table_html, 'lxml')
					table = soup.find('table')
					rows = table.find_all('tr')
					# 创建Word文档
					doc = Document()
					# 创建Word表格
					word_table = doc.add_table(rows=0, cols=len(rows[0].find_all(['td', 'th'])))
					# 定义边框样式（实线、1pt宽度、黑色）
					border_style = {
						'val': 'single',  # 边框类型：单线
						'sz': 6,  # 边框宽度（8代表1pt）
						'color': '000000'  # 黑色
					}
					# 获取表格属性对象
					tbl = word_table._tbl
					tbl_properties = tbl.tblPr
					# 设置表格边框
					tbl_borders = OxmlElement('w:tblBorders')
					for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
						border = OxmlElement(f'w:{border_name}')
						for key, value in border_style.items():
							border.set(qn(f'w:{key}'), str(value))
						tbl_borders.append(border)
					# 将边框属性添加到表格属性中
					tbl_properties.append(tbl_borders)
					# 将HTML表格数据写入Word表格
					for row in rows:
						cells = row.find_all(['td', 'th'])
						word_row = word_table.add_row().cells
						for i, cell in enumerate(cells):
							word_row[i].text = cell.get_text()
					# 调整表格列宽（可选）
					for column in word_table.columns:
						column.width = 2000000  # 设置列宽为2厘米（2000000 twips）

					# 保存Word文档
					folder_path = r".\output_word"
					os.makedirs(folder_path, exist_ok=True)
					full_path = os.path.join(folder_path, file_name)
					doc.save(full_path)
					return True
			except Exception as e:
				return False
		else:
			return False

	@staticmethod
	def convert_to_html(file_path, mode, file_name="to_html.html"):
		if mode == 1:
			doc = Document(file_path)
			table = doc.tables[0]
			list_cells = []
			for row_idx, row in enumerate(table.rows):
				row_cells = []  # 存储当前行的所有单元格对象
				for cell_idx, cell in enumerate(row.cells):
					cells = Convertor.check_cell_style(cell)
					cells.text = cell.text.strip()  # 设置单元格的文本内容（去除首尾空格）
					# 记录单元格的行列位置
					cells.row = row_idx + 1
					cells.col = cell_idx + 1
					# cells.colSpan = is_grid_span(cell)
					row_cells.append(cells)  # 所有行的单元格对象
				# 存储提取到的文本和该行单元格信息
				list_cells.append(row_cells)

			html_content = Convertor.convertToHtml(list_cells)
			folder_path = r".\output_html"
			os.makedirs(folder_path, exist_ok=True)
			full_path = os.path.join(folder_path, file_name)
			with open(full_path, "w", encoding="utf-8") as f:
				f.write(html_content)
				return True
		elif mode == 2:
			return Convertor.pdf_to_html(file_path, file_name=file_name)
		else:
			return False

	@staticmethod
	def pdf_to_html(file_path, file_name="pdf_to_html.html"):
		with pdfplumber.open(file_path) as pdf:
			first_page = pdf.pages[0]
			table = first_page.extract_table()
			# 将表格数据转换为HTML表格格式
			html_template = f"""<!DOCTYPE html>
			<html>
			<head>
			    <meta name="viewport" content="width=device-width, initial-scale=1.0">
			    <style>
			        .data-table {{
			            width: 80%;
			            margin: 20px auto;
			            border-collapse: collapse;
			            font-family: Arial, sans-serif;
			            box-shadow: 0 1px 3px rgba(0,0,0,0.2);
			        }}
			        .data-table td, .data-table th {{
			            padding: 12px 15px;
			            text-align: center;
			            border: 1px solid #ddd;
			        }}
			        .data-table tbody tr:nth-child(even) {{
			            background-color: #f9f9f9;
			        }}
			        .data-table tbody tr:hover {{
			            background-color: #f1f1f1;
			            transition: background-color 0.3s;
			        }}
			        .table-container {{
			            width: 100%;
			            overflow-x: auto;
			        }}
			    </style>
			</head>
			<body>
			    <div class="table-container">
			        <table class="data-table">
			            <tbody>
			                {''.join(
				f'<tr>{"".join(f"<td>{html.escape(str(cell))}</td>" for cell in row)}</tr>'
				for row in table
			)}
			            </tbody>
			        </table>
			    </div>
			</body>
			</html>
			"""
			try:
				folder_path = r".\output_html"
				os.makedirs(folder_path, exist_ok=True)
				full_path = os.path.join(folder_path, file_name)
				with open(full_path, "w", encoding="utf-8") as html_file:
					html_file.write(html_template)
				return True
			except  Exception as e:
				return False

	@staticmethod
	def convert_to_pdf_micro(file_path, file_name="word_to_pdf.pdf"):
		try:
			folder_path = r".\output_pdf"
			os.makedirs(folder_path, exist_ok=True)
			full_path = os.path.join(folder_path, file_name)
			convert(file_path, full_path)
			return True
		except Exception as e:
			return False

